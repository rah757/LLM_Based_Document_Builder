"""
Final document generation utilities
Handles replacement of placeholder markers with actual values
"""

from docx import Document
import re


def replace_placeholders_in_document(marked_doc_path, placeholders, output_path):
    """
    Replace placeholder markers with actual values in the marked document
    
    Args:
        marked_doc_path: Path to the placeholder-marked document
        placeholders: List of placeholder dictionaries with user_input
        output_path: Path to save the final document
    
    Returns:
        dict: {'success': bool, 'replacements': int, 'errors': list}
    """
    try:
        doc = Document(marked_doc_path)
        
        # Build replacement map: [placeholder_id: description] -> user_input
        replacement_map = {}
        for p in placeholders:
            placeholder_id = p['placeholder_id']
            description = p['description']
            user_input = p.get('user_input')
            
            # The marked document has patterns like: [placeholder_001: the 'Company Name']
            marker_pattern = f"[{placeholder_id}: {description}]"
            
            if user_input:
                replacement_map[marker_pattern] = user_input
        
        replacements_made = 0
        errors = []
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text
            
            for marker, value in replacement_map.items():
                if marker in paragraph_text:
                    # Replace in runs to preserve formatting
                    replaced = False
                    
                    # Try to find and replace within runs
                    for run in paragraph.runs:
                        if marker in run.text:
                            run.text = run.text.replace(marker, value)
                            replacements_made += 1
                            replaced = True
                    
                    # If not found in individual runs, might be split across runs
                    if not replaced and marker in paragraph_text:
                        # Rebuild paragraph text
                        new_text = paragraph_text.replace(marker, value)
                        # Clear existing runs
                        for run in paragraph.runs:
                            run.text = ''
                        # Add new text to first run
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                        else:
                            paragraph.add_run(new_text)
                        replacements_made += 1
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_text = paragraph.text
                        
                        for marker, value in replacement_map.items():
                            if marker in paragraph_text:
                                # Replace in runs
                                replaced = False
                                
                                for run in paragraph.runs:
                                    if marker in run.text:
                                        run.text = run.text.replace(marker, value)
                                        replacements_made += 1
                                        replaced = True
                                
                                # If not found in individual runs
                                if not replaced and marker in paragraph_text:
                                    new_text = paragraph_text.replace(marker, value)
                                    for run in paragraph.runs:
                                        run.text = ''
                                    if paragraph.runs:
                                        paragraph.runs[0].text = new_text
                                    else:
                                        paragraph.add_run(new_text)
                                    replacements_made += 1
        
        # Save the final document
        doc.save(output_path)
        
        return {
            'success': True,
            'replacements': replacements_made,
            'errors': errors
        }
        
    except Exception as e:
        return {
            'success': False,
            'replacements': 0,
            'errors': [str(e)]
        }


def check_all_filled(placeholders):
    """
    Check if all placeholders are filled (or auto_filled)
    
    Args:
        placeholders: List of placeholder dictionaries
    
    Returns:
        dict: {'all_filled': bool, 'pending': list, 'has_auto_filled': bool}
    """
    pending = []
    has_auto_filled = False
    
    for p in placeholders:
        status = p.get('status', 'pending')
        
        if status == 'pending':
            pending.append(p['placeholder_id'])
        elif status == 'auto_filled':
            has_auto_filled = True
    
    return {
        'all_filled': len(pending) == 0,
        'pending': pending,
        'has_auto_filled': has_auto_filled
    }

