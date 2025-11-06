"""
Placeholder detection and processing utilities
Handles placeholder detection, context extraction, and document marking
"""

import re
import json
from docx import Document


def extract_document_text(doc):
    """
    Extract full text from a docx Document object
    
    Args:
        doc: python-docx Document object
    
    Returns:
        str: Extracted text
    """
    full_text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    full_text.append(cell.text)
    
    return '\n'.join(full_text)


def extract_context(text, start_pos, end_pos, words_count):
    """
    Extract context before and after a placeholder position
    
    Args:
        text: The full document text
        start_pos: Start position of the placeholder
        end_pos: End position of the placeholder
        words_count: Number of words to extract before and after
    
    Returns:
        tuple: (context_before, context_after, before_words_actual, after_words_actual)
    """
    # Extract text before the placeholder
    text_before = text[:start_pos].strip()
    words_before = text_before.split()
    context_before = ' '.join(words_before[-words_count:]) if len(words_before) > words_count else text_before
    before_words_actual = min(len(words_before), words_count)
    
    # Extract text after the placeholder
    text_after = text[end_pos:].strip()
    words_after = text_after.split()
    context_after = ' '.join(words_after[:words_count]) if len(words_after) > words_count else text_after
    after_words_actual = min(len(words_after), words_count)
    
    return context_before, context_after, before_words_actual, after_words_actual


def detect_placeholders(text, patterns_config):
    """
    Detect placeholders in the document text using regex patterns from config
    
    Args:
        text: The document text
        patterns_config: Configuration dict with patterns
    
    Returns:
        list: List of dictionaries containing placeholder information
    """
    patterns = patterns_config.get('patterns', [])
    context_words_count = patterns_config.get('context_words_count', 20)
    
    # Filter enabled patterns only
    enabled_patterns = [p for p in patterns if p.get('enabled', True)]
    
    placeholders = []
    placeholder_counter = 1
    
    for pattern_config in enabled_patterns:
        pattern = pattern_config.get('regex')
        pattern_name = pattern_config.get('name', 'unknown')
        
        if not pattern:
            continue
        
        try:
            for match in re.finditer(pattern, text):
                placeholder_text = match.group(0)
                placeholder_name = match.group(1).strip()
                start_pos = match.start()
                end_pos = match.end()
                
                # Extract context using configured word count
                context_before, context_after, before_words_actual, after_words_actual = extract_context(
                    text, start_pos, end_pos, context_words_count
                )
                
                placeholder_info = {
                    "placeholder_id": f"placeholder_{str(placeholder_counter).zfill(3)}",
                    "placeholder": placeholder_text,
                    "placeholder_name": placeholder_name,
                    "description": f"the '{placeholder_name}'",
                    "pattern_type": pattern_name,
                    "position": {
                        "start": start_pos,
                        "end": end_pos
                    },
                    "context_before": context_before,
                    "context_after": context_after,
                    "context_window": {
                        "before_words": before_words_actual,
                        "after_words": after_words_actual
                    },
                    "user_input": None,
                    "status": "pending"
                }
                
                placeholders.append(placeholder_info)
                placeholder_counter += 1
        except re.error as e:
            print(f"Error in regex pattern '{pattern_name}': {str(e)}")
            continue
    
    # Sort placeholders by position to maintain order
    placeholders.sort(key=lambda x: x['position']['start'])
    
    return placeholders


def create_marked_document(original_doc_path, placeholders, output_path):
    """
    Create a new document with placeholders marked with unique IDs
    Handles duplicate placeholders by replacing each occurrence sequentially
    
    Args:
        original_doc_path: Path to the original document
        placeholders: List of placeholder dictionaries (sorted by position)
        output_path: Path where the marked document will be saved
    
    Returns:
        str: Path to the saved marked document
    """
    doc = Document(original_doc_path)
    
    # Process each placeholder ONE AT A TIME (handles duplicates)
    # Sort by position to ensure we replace in document order
    placeholders_sorted = sorted(placeholders, key=lambda x: x['position']['start'])
    
    for ph in placeholders_sorted:
        original = ph['placeholder']
        marked = f"[{ph['placeholder_id']}: {ph['description']}]"
        replaced = False
        
        # Replace the FIRST occurrence found
        # Try paragraphs first - check paragraph.text (handles split runs)
        for paragraph in doc.paragraphs:
            if replaced:
                break
            if original in paragraph.text:
                # Replace in paragraph text (works even if split across runs)
                new_text = paragraph.text.replace(original, marked, 1)  # Replace FIRST occurrence
                
                # Clear all runs and rebuild
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)
                
                replaced = True
                break
    
        # If not replaced in paragraphs, try tables
        if not replaced:
            for table in doc.tables:
                if replaced:
                    break
                for row in table.rows:
                    if replaced:
                        break
                    for cell in row.cells:
                        if replaced:
                            break
                        for paragraph in cell.paragraphs:
                            if replaced:
                                break
                            if original in paragraph.text:
                                # Replace in paragraph text (handles split runs)
                                new_text = paragraph.text.replace(original, marked, 1)
                                
                                # Clear all runs and rebuild
                                for run in paragraph.runs:
                                    run.text = ''
                                if paragraph.runs:
                                    paragraph.runs[0].text = new_text
                                else:
                                    paragraph.add_run(new_text)
                                
                                replaced = True
                                break
        
        if not replaced:
            print(f"Warning: Could not find '{original}' for {ph['placeholder_id']}")
    
    # Save the marked document
    doc.save(output_path)
    return output_path

